#!/usr/bin/env python3
# coding: utf-8

"""
Creatore di Digest Bot
- Принимает Excel-файл (имя канала | адрес/username)
- Запрашивает интервал: сутки/неделя/месяц/произвольный
- Запрашивает ключевые слова / теги
- Скачивает сообщения из каналов (через Telethon или Bot API если настроен)
- Фильтрует по ключам, суммаризует (NLTK), формирует .docx
- Поддерживает опцию автодоставки по расписанию (apscheduler)
- Кеширует скачанные сообщения в sqlite
"""

import os
import logging
import tempfile
import sqlite3
import asyncio
import html
from datetime import datetime, timedelta
from io import BytesIO
from typing import List, Tuple, Optional

import pandas as pd
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem.snowball import SnowballStemmer
from docx import Document

from apscheduler.schedulers.asyncio import AsyncIOScheduler

from telethon import TelegramClient, errors as telethon_errors
from telethon.tl.types import PeerChannel

from telegram import (
    Update, 
    InlineKeyboardButton, 
    InlineKeyboardMarkup, 
    ReplyKeyboardRemove,
    KeyboardButton, 
    ReplyKeyboardMarkup, 
    InputFile
)
from telegram.ext import (
    Application,
    ApplicationBuilder, 
    CommandHandler, 
    MessageHandler, 
    CallbackQueryHandler, 
    ContextTypes,
    filters,
    ConversationHandler
)

# -----------------------------
# Конфиг и логирование
# -----------------------------
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# -----------------------------
# Состояния
# -----------------------------
WAITING_FOR_FILE = 1
WAITING_FOR_INTERVAL = 2
WAITING_FOR_CUSTOM_INTERVAL_FROM = 3
WAITING_FOR_CUSTOM_INTERVAL_TO = 4
WAITING_FOR_KEYWORDS = 5

# NLTK setup (будет скачивать при первом запуске)
nltk_resources = ["punkt", "stopwords"]
for res in nltk_resources:
    try:
        nltk.data.find(res)
    except LookupError:
        nltk.download(res)

STOPWORDS = set(stopwords.words("russian")) | set(stopwords.words("english"))
STEMMER = SnowballStemmer("russian")

# -----------------------------
# Безопасное получение секретов
# -----------------------------
def get_secret(name: str, docker_secret_path: Optional[str] = None) -> Optional[str]:
    """
    Получение секрета:
    - Сначала пытаемcя прочитать Docker secret (по пути /run/secrets/<name>), если существует
    - Иначе берем из переменных окружения
    """
    # Docker secrets location (standard)
    secret_file = f"/run/secrets/{name}"
    if os.path.exists(secret_file):
        with open(secret_file, "r") as f:
            return f.read().strip()
    # fallback to env
    return os.getenv(name)

def get_telegram_token() -> str:
    token = get_secret("TELEGRAM_BOT_TOKEN")
    if not token:
        raise RuntimeError("TELEGRAM_BOT_TOKEN не задан. Установите в окружении или Docker Secret.")
    return token

def get_telethon_credentials() -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """
    Возвращает (api_id, api_hash, session_string).
    Если api_id/api_hash заданы, можно использовать Telethon для чтения каналов.
    Если session_string заданы — Telethon может использовать их.
    """
    api_id = get_secret("TELETHON_API_ID")
    api_hash = get_secret("TELETHON_API_HASH")
    session = get_secret("TELETHON_SESSION")  # опционально
    return api_id, api_hash, session

# Глобальная переменная для клиента
client = TelegramClient('session_name', int(API_ID), API_HASH)

# -----------------------------
# Простой SQLite кеш для сообщений
# -----------------------------
# Путь к базе берём из переменной окружения или используем дефолт
DB_PATH = os.getenv("DB_PATH", "/app/data/digest_cache.sqlite")

def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)  # создаём папку если её нет
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS digests (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            channel_name TEXT,
            channel_link TEXT,
            content TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

def cache_posts(posts: List[Tuple[str,str,str,str]]):
    """
    posts: list of tuples (id, channel, date_iso, text)
    """
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    for pid, channel, date_iso, text in posts:
        try:
            cur.execute("INSERT OR IGNORE INTO posts (id, channel, date, text) VALUES (?, ?, ?, ?)",
                        (pid, channel, date_iso, text))
        except Exception as e:
            logger.exception("DB insert failed: %s", e)
    conn.commit()
    conn.close()

def query_posts(channel: str, date_from: datetime, date_to: datetime) -> List[Tuple[str,str,str,str]]:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, channel, date, text FROM posts WHERE channel=? AND date BETWEEN ? AND ?",
                (channel, date_from.isoformat(), date_to.isoformat()))
    rows = cur.fetchall()
    conn.close()
    return rows

# -----------------------------
# Утилиты для суммаризации (на базе NLTK)
# -----------------------------
def normalize_word(w: str) -> str:
    w = w.lower()
    w = ''.join(ch for ch in w if ch.isalpha() or ch == '-')
    if not w: 
        return w
    try:
        return STEMMER.stem(w)
    except Exception:
        return w

def text_score_sentences(text: str, keywords: List[str], top_n_sentences: int = 3) -> List[str]:
    """
    Простая суммаризация: среди предложений выбираем top_n по совпадениям с ключевыми словами
    или по частоте важных слов.
    """
    if not text or text.strip() == "":
        return []
    sents = sent_tokenize(text)
    if not sents:
        return []
    # build keyword set
    kw_norm = set(normalize_word(k) for k in keywords if k.strip())
    # score each sentence
    scores = []
    for s in sents:
        words = [normalize_word(w) for w in word_tokenize(s)]
        # keyword hits
        hits = sum(1 for w in words if w in kw_norm)
        # informative score: number of non-stopword tokens
        informative = sum(1 for w in words if w and w not in STOPWORDS)
        total = hits * 5 + informative  # bias keywords
        scores.append((total, s))
    # sort
    scores.sort(key=lambda x: x[0], reverse=True)
    chosen = [s for _, s in scores[:top_n_sentences]]
    # If keywords empty and no hits, fall back to first N sentences
    if not chosen:
        chosen = sents[:top_n_sentences]
    # deduplicate and preserve order appearing in original
    chosen_set = set(chosen)
    final = [s for s in sents if s in chosen_set]
    return final[:top_n_sentences]

# -----------------------------
# Telethon client helper
# -----------------------------
async def fetch_messages_telethon(api_id: str, api_hash: str, session: Optional[str], channel_username: str,
                                  date_from: datetime, date_to: datetime, limit: int = 200) -> List[Tuple[str,str,str,str]]:
    """
    Получить сообщения из канала через Telethon (user client).
    Возвращает список кортежей (id, channel_username, date_iso, text)
    """
    # Session: если передана строка — можно использовать ее как session name
    session_name = session or "digest_bot_session"
    client = TelegramClient(session_name, int(api_id), api_hash)
    await client.start()
    posts = []
    try:
        # resolve entity
        try:
            entity = await client.get_entity(channel_username)
        except Exception:
            # last resort: try to get by username as PeerChannel (some public channels)
            raise

        # iterate messages in date range
        async for msg in client.iter_messages(entity, limit=limit, reverse=False):
            if not msg.date:
                continue
            if msg.date.replace(tzinfo=None) < date_from:
                continue
            if msg.date.replace(tzinfo=None) > date_to:
                continue
            text = msg.message or ""
            pid = f"{entity.id}_{msg.id}"
            posts.append((pid, channel_username, msg.date.isoformat(), text))
    except telethon_errors.rpcerrorlist.ChannelPrivateError:
        logger.warning("Channel %s is private or access denied.", channel_username)
    except Exception as e:
        logger.exception("Telethon fetch error for %s: %s", channel_username, e)
    finally:
        await client.disconnect()
    return posts

# -----------------------------
# Фоллбек: получить через Bot API (если бот админ канала)
# -----------------------------
# NOTE: Telegram Bot API generally не позволяет получать историю публичных каналов,
# если бот не добавлен и не имеет прав. Этот метод может работать только когда бот является администратором/имеет доступ.
# В python-telegram-bot есть метод get_chat и get_chat_history в форме get_updates? В v20 Bot API не предоставляет прямой fetch history,
# поэтому мы оставим это как stub / placeholder и рекомендуем использовать Telethon.
async def fetch_messages_botapi_stub(application, chat_identifier: str, date_from: datetime, date_to: datetime):
    # stub: вернуть пустой
    logger.info("Bot API fetch stub called for %s — рекомендуется настроить Telethon (user API) или сделать бота админом каналов.", chat_identifier)
    return []

# -----------------------------
# Создание .docx с дайджестом
# -----------------------------
def create_docx_digest(digest_entries: List[Tuple[str, str, str, List[str]]], title: str = "Дайджест") -> bytes:
    """
    digest_entries: list of tuples (channel_name, channel_link, post_date_iso, summary_sentences_list)
    Возвращает bytes docx
    """
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Сгенерировано: {datetime.utcnow().isoformat()} UTC")
    for channel, link, date_iso, summary_sentences in digest_entries:
        p = doc.add_paragraph()
        p.add_run(channel).bold = True
        if link:
            p.add_run(f" — {link}")
        p.add_run(f" ({date_iso})\n")
        for sent in summary_sentences:
            doc.add_paragraph(sent, style='List Bullet')
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# -----------------------------
# Логика бота / Handlers
# -----------------------------
# We'll keep a simple in-memory per-chat state (single-user assumption)
CHAT_STATE = {}

def reset_chat_state(chat_id: int):
    CHAT_STATE[chat_id] = {
        "excel_path": None,
        "channels_df": None,
        "interval": None,
        "date_from": None,
        "date_to": None,
        "keywords": [],
        "use_telethon": False
    }

# Helpers to parse interval choice
def parse_interval_choice(choice: str) -> Tuple[datetime, datetime]:
    now = datetime.utcnow()
    if choice == "Сутки":
        return now - timedelta(days=1), now
    elif choice == "Неделя":
        return now - timedelta(weeks=1), now
    elif choice == "Месяц":
        return now - timedelta(days=30), now
    else:
        raise ValueError("unsupported quick choice")

# Build interval keyboard
def interval_keyboard():
    keyboard = [
        [InlineKeyboardButton("Сутки", callback_data="interval_Сутки"),
         InlineKeyboardButton("Неделя", callback_data="interval_Неделя")],
        [InlineKeyboardButton("Месяц", callback_data="interval_Месяц"),
         InlineKeyboardButton("Задайте произвольный интервал", callback_data="interval_custom")]
    ]
    return InlineKeyboardMarkup(keyboard)

# Start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    reset_chat_state(chat_id)
    await update.message.reply_text(
        "Привет! Отправь мне Excel-файл (xlsx) со списком каналов (столбцы: имя канала, адрес/username).\n"
        "Файл должен содержать 2 столбца: 'name' и 'address' (или первые два столбца будут использованы)."
    )
    return WAITING_FOR_FILE

# -----------------------------
# Обработка загруженного файла
# -----------------------------
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = update.message.document

    if not document:
        await update.message.reply_text("Пожалуйста, загрузите Excel-файл.")
        return WAITING_FOR_FILE

    with tempfile.NamedTemporaryFile(delete=False) as tf:
        tg_file = await document.get_file()
        await tg_file.download_to_drive(tf.name)
        file_path = tf.name

    ext = os.path.splitext(document.file_name)[-1].lower()
    try:
        if ext == ".xlsx":
            df = pd.read_excel(file_path, engine="openpyxl")
        elif ext == ".xls":
            df = pd.read_excel(file_path, engine="xlrd")
        else:
            await update.message.reply_text("Неподдерживаемый формат. Используйте .xls или .xlsx")
            return WAITING_FOR_FILE
    except Exception as e:
        await update.message.reply_text(f"Ошибка при чтении Excel: {e}")
        return WAITING_FOR_FILE
    
    # Сохраняем таблицу каналов для текущего пользователя
    context.user_data["channels"] = df

    # Отправляем кнопки выбора интервала
    keyboard = [
        [InlineKeyboardButton("Сутки", callback_data="interval_day")],
        [InlineKeyboardButton("Неделя", callback_data="interval_week")],
        [InlineKeyboardButton("Месяц", callback_data="interval_month")],
        [InlineKeyboardButton("Задайте произвольный интервал", callback_data="interval_custom")]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Выберите интервал времени:", reply_markup=reply_markup)

    return WAITING_FOR_INTERVAL

async def handle_interval(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    if text in ["Сутки", "Неделя", "Месяц"]:
        context.user_data["interval"] = text
        await update.message.reply_text(
            f"⏳ Вы выбрали интервал: {text}\n\nТеперь введите ключевые слова или теги (через запятую):",
            reply_markup=ReplyKeyboardRemove()
        )
    elif text == "Задайте произвольный интервал":
        context.user_data["interval"] = "custom"
        await update.message.reply_text(
            "✍️ Введите произвольный интервал в формате: `YYYY-MM-DD до YYYY-MM-DD`",
            parse_mode="Markdown",
            reply_markup=ReplyKeyboardRemove()
        )
    else:
        await update.message.reply_text("⚠️ Пожалуйста, выберите один из предложенных интервалов.")

# -----------------------------
# Callback при выборе интервала
# -----------------------------
async def interval_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    data = query.data.replace("interval_", "")

    if data == "custom":
        await query.edit_message_text("Введите дату начала интервала (ГГГГ-ММ-ДД):")
        return WAITING_FOR_CUSTOM_INTERVAL_FROM
    else:
        # Сохраняем выбранный интервал
        context.user_data["interval"] = data
        await query.edit_message_text("Введите ключевые слова (через запятую):")
        return WAITING_FOR_KEYWORDS

# -----------------------------
# Кастомный интервал: начало
# -----------------------------
async def custom_interval_from(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["custom_from"] = update.message.text.strip()
    await update.message.reply_text("Введите дату окончания интервала (ГГГГ-ММ-ДД):")
    return WAITING_FOR_CUSTOM_INTERVAL_TO

# -----------------------------
# Кастомный интервал: конец
# -----------------------------
async def custom_interval_to(update: Update, context: ContextTypes.DEFAULT_TYPE):
    custom_from = context.user_data.get("custom_from")
    custom_to = update.message.text.strip()
    context.user_data["interval"] = (custom_from, custom_to)

    await update.message.reply_text("Введите ключевые слова (через запятую):")
    return WAITING_FOR_KEYWORDS

# -----------------------------
# Обработка ключевых слов
# -----------------------------
async def handle_keywords(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keywords = [k.strip() for k in update.message.text.split(",") if k.strip()]
    context.user_data["keywords"] = keywords

    await update.message.reply_text(
        "Файл принят ✅\nИнтервал задан ✅\nКлючевые слова сохранены ✅\n\nГотовлю дайджест...",
        reply_markup=ReplyKeyboardRemove()
    )

    # Асинхронно вызываем generate_digest
    digest_path = await generate_digest(context.user_data)  # Используем await здесь

    if digest_path and os.path.exists(digest_path):
        await update.message.reply_document(open(digest_path, "rb"), filename="digest.docx")
    else:
        await update.message.reply_text("Не удалось создать дайджест 😢")

    return ConversationHandler.END

async def get_posts(channel_link, interval):
    await client.start()
    channel = await client.get_entity(channel_link)
    now = datetime.utcnow()
    
    # Определяем дату начала интервала
    if interval == "day":
        start_date = now - timedelta(days=1)
    elif interval == "week":
        start_date = now - timedelta(weeks=1)
    elif interval == "month":
        start_date = now - timedelta(days=30)
    elif isinstance(interval, tuple):
        start_date = datetime.fromisoformat(interval[0])
        end_date = datetime.fromisoformat(interval[1])
    else:
        start_date = now - timedelta(days=1)
    end_date = now if not isinstance(interval, tuple) else end_date

    posts_text = []
    async for message in client.iter_messages(channel, offset_date=end_date, reverse=True):
        if message.date < start_date:
            break
        if message.text:
            posts_text.append((message.date, message.text))
        else:
            print(f"Message without text found: {message.date}")
    
    print(f"Found {len(posts_text)} posts in channel: {channel_link}")
    return posts_text

async def generate_digest(user_data):
    channels = user_data.get("channels")
    interval = user_data.get("interval")
    keywords = user_data.get("keywords", [])

    if channels is None or not keywords:
        return None

    await client.start()

    digest_text = "📌 Дайджест по вашим каналам:\n\n"

    for _, row in channels.iterrows():
        channel_name = row[0]
        channel_link = row[1]

        # Получаем посты
        posts = await get_posts(channel_link, interval)
        if not posts:
            digest_text += f"{channel_name} ({channel_link}): Нет сообщений за этот интервал\n"
            continue
        
        digest_text += f"--- {channel_name} ({channel_link}) ---\n"
        for date, text in posts:
            if text:  # Проверяем, что текст не пустой
                summary = summarize_text(text, keywords)
                digest_text += f"{date.date()}: {summary}\n"
            else:
                digest_text += f"{date.date()}: (Пустое сообщение)\n"

    output_dir = "/app/data"
    os.makedirs(output_dir, exist_ok=True)
    digest_path = os.path.join(output_dir, "digest.docx")

    doc = Document()
    doc.add_heading("Дайджест", 0)
    doc.add_paragraph(digest_text)
    doc.save(digest_path)

    return digest_path

# Core processing
async def process_digest_for_chat(chat_id: int, context: ContextTypes.DEFAULT_TYPE):
    state = CHAT_STATE.get(chat_id)
    if not state:
        await context.bot.send_message(chat_id, "Внутренняя ошибка: состояние чата не найдено. Запустите /start.")
        return

    df: pd.DataFrame = state.get("channels_df")
    date_from: datetime = state.get("date_from")
    date_to: datetime = state.get("date_to")
    keywords: List[str] = state.get("keywords", [])

    if df is None or date_from is None or date_to is None:
        await context.bot.send_message(chat_id, "Нечего обрабатывать — отсутствуют данные. Запустите заново /start.")
        return

    await context.bot.send_message(chat_id, f"Собираю посты из {len(df)} каналов за период {date_from.date()} — {date_to.date()}...")

    # choose Telethon if configured
    api_id, api_hash, session = get_telethon_credentials()
    use_telethon = False
    if api_id and api_hash:
        use_telethon = True

    collected_posts = []  # list of (id, channel, date_iso, text)
    # iterate channels
    for idx, row in df.iterrows():
        chan_name = str(row['name'])
        chan_addr = str(row['address'])
        try:
            if use_telethon:
                try:
                    posts = await fetch_messages_telethon(api_id, api_hash, session, chan_addr, date_from, date_to, limit=500)
                except Exception as e:
                    logger.exception("Telethon per-channel error: %s", e)
                    posts = []
            else:
                posts = await fetch_messages_botapi_stub(context.application, chan_addr, date_from, date_to)
            # if posts empty - try cached posts
            if not posts:
                cached = query_posts(chan_addr, date_from, date_to)
                if cached:
                    posts = cached
            if posts:
                collected_posts.extend(posts)
                cache_posts(posts)
        except Exception as e:
            logger.exception("Error fetching channel %s: %s", chan_addr, e)

    if not collected_posts:
        await context.bot.send_message(chat_id, "За указанный период постов не найдено.")
        return

    await context.bot.send_message(chat_id, f"Найдено {len(collected_posts)} постов. Выполняю суммаризацию...")

    # Group posts by channel and date, make summaries
    digest_entries = []
    grouped = {}
    for pid, channel, date_iso, text in collected_posts:
        key = (channel, date_iso[:10])  # group by date (YYYY-MM-DD)
        grouped.setdefault(key, []).append(text or "")

    # For each group produce summary
    for (channel, date_only), texts in grouped.items():
        combined_text = "\n".join(texts)
        # our summarizer: get top sentences
        top_sents = text_score_sentences(combined_text, keywords, top_n_sentences=4)
        # create link if channel looks like @username or t.me/...
        link = None
        if channel.startswith("@"):
            link = f"https://t.me/{channel[1:]}"
        elif channel.startswith("http"):
            link = channel
        else:
            # try direct t.me
            link = f"https://t.me/{channel}"
        digest_entries.append((channel, link, date_only, top_sents))

    # Create docx
    docx_bytes = create_docx_digest(digest_entries, title=f"Дайджест {date_from.date()} — {date_to.date()}")
    await context.bot.send_document(chat_id, document=InputFile(BytesIO(docx_bytes), filename=f"digest_{date_from.date()}_{date_to.date()}.docx"))
    await context.bot.send_message(chat_id, "Дайджест готов и отправлен. Спасибо!")

# Command to schedule regular digest (optional improvement)
async def schedule_digest_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    # This is a simple example that schedules every day at 09:00 UTC
    await update.message.reply_text("Запрошено ежедневное расписание дайджеста в 09:00 UTC. (Пример)")
    scheduler = context.bot_data.get("scheduler")
    if not scheduler:
        scheduler = AsyncIOScheduler()
        scheduler.start()
        context.bot_data["scheduler"] = scheduler

    async def job_fn():
        # For simplicity, use existing chat state
        await process_digest_for_chat(chat_id, context)

    scheduler.add_job(job_fn, 'cron', hour=9, minute=0, id=f"digest_{chat_id}", replace_existing=True)
    await update.message.reply_text("Ежедневная отправка запланирована (09:00 UTC).")

# Cancel
async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat_id = update.effective_chat.id
    reset_chat_state(chat_id)
    await update.message.reply_text("Операция отменена. Если хотите, начните заново /start.")
    return ConversationHandler.END

# Fallback unknown messages
async def unknown(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Не понял. Используйте /start чтобы начать или отправьте файл Excel с каналами.")

# -----------------------------
# Main
# -----------------------------
def main():
    # Инициализация БД
    init_db()

    # Получение токена бота безопасно
    token = get_telegram_token()

    # Создаем приложение
    application = ApplicationBuilder().token(token).build()

    # ConversationHandler для сценария загрузки файла и дайджеста
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],

        states={
            WAITING_FOR_FILE: [
                MessageHandler(filters.Document.ALL, handle_file)
            ],
            WAITING_FOR_INTERVAL: [
                CallbackQueryHandler(interval_callback, pattern=r"^interval_")
            ],
            WAITING_FOR_CUSTOM_INTERVAL_FROM: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, custom_interval_from)
            ],
            WAITING_FOR_CUSTOM_INTERVAL_TO: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, custom_interval_to)
            ],
            WAITING_FOR_KEYWORDS: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_keywords)
            ]
        },

        fallbacks=[CommandHandler("cancel", cancel)]
    )

    # Регистрация хендлеров
    application.add_handler(conv_handler)
    application.add_handler(CommandHandler("schedule", schedule_digest_cmd))
    application.add_handler(MessageHandler(filters.COMMAND, unknown))  # неизвестные команды

    # Старт бота
    logger.info("🤖 Бот запущен...")
    application.run_polling()


# Точка входа
if __name__ == "__main__":
    main()